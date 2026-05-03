#!/usr/bin/env python3
"""
md → docx 後處理：修正 pandoc 產出的表格欄寬問題
用法：python3 fix_tables.py <docx路徑>
"""
import sys
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


BORDER = {"w:val": "single", "w:sz": "6", "w:space": "0", "w:color": "auto"}


def make_border_el(tag):
    el = OxmlElement(tag)
    for k, v in BORDER.items():
        el.set(qn(k), v)
    return el


def add_table_borders(tblPr):
    old = tblPr.find(qn("w:tblBorders"))
    if old is not None:
        tblPr.remove(old)
    tblBorders = OxmlElement("w:tblBorders")
    for tag in ("w:top", "w:left", "w:bottom", "w:right", "w:insideH", "w:insideV"):
        tblBorders.append(make_border_el(tag))
    tblPr.append(tblBorders)


def add_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    old = tcPr.find(qn("w:tcBorders"))
    if old is not None:
        tcPr.remove(old)
    tcBorders = OxmlElement("w:tcBorders")
    for tag in ("w:top", "w:left", "w:bottom", "w:right"):
        tcBorders.append(make_border_el(tag))
    tcPr.append(tcBorders)


def fix_tables(path):
    doc = Document(path)

    for table in doc.tables:
        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)

        tblW = tblPr.find(qn("w:tblW"))
        if tblW is None:
            tblW = OxmlElement("w:tblW")
            tblPr.append(tblW)
        tblW.set(qn("w:w"), "5000")
        tblW.set(qn("w:type"), "pct")

        tblLayout = tblPr.find(qn("w:tblLayout"))
        if tblLayout is None:
            tblLayout = OxmlElement("w:tblLayout")
            tblPr.append(tblLayout)
        tblLayout.set(qn("w:type"), "autofit")

        add_table_borders(tblPr)

        for row in table.rows:
            for cell in row.cells:
                tcPr = cell._tc.find(qn("w:tcPr"))
                if tcPr is not None:
                    tcW = tcPr.find(qn("w:tcW"))
                    if tcW is not None:
                        tcW.set(qn("w:w"), "0")
                        tcW.set(qn("w:type"), "auto")
                add_cell_borders(cell)

    doc.save(path)
    print(f"Done: {path}（{len(doc.tables)} 張表格已修正）")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法：python3 fix_tables.py <docx路徑>")
        sys.exit(1)
    fix_tables(sys.argv[1])
